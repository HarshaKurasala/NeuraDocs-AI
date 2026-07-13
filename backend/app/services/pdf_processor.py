"""
services/pdf_processor.py - PDF text extraction with multi-library fallback.

Strategy:
  1. PyMuPDF (fitz)  — fastest, best layout preservation
  2. pdfplumber      — better for tables/complex layouts
  3. pypdf           — pure-Python fallback

Each page's text is returned with its metadata (page number, source file)
so citations can reference exact page numbers later.
"""

import fitz  # PyMuPDF
import pdfplumber
from pypdf import PdfReader
from docx import Document as DocxDocument
from pathlib import Path
from typing import Optional

from app.utils.logger import get_logger

logger = get_logger(__name__)


TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".csv"}


def _normalize_text(text: str) -> str:
    """Normalize extracted text while preserving whitespace.

    Problem observed:
      - aggressive normalization collapses formatting in ways that can remove
        the spaces between words when rendered in the UI.

    This implementation:
      - trims only leading/trailing whitespace per extracted line
      - keeps internal spaces intact
      - keeps line breaks between extracted lines/paragraphs
      - removes only fully empty lines
    """

    if not text:
        return ""

    cleaned_lines: list[str] = []
    for line in text.splitlines():
        normalized = line.strip()  # trim only edges; preserve internal spacing
        if normalized:
            cleaned_lines.append(normalized)

    return "\n".join(cleaned_lines).strip()


def extract_text_docx(file_path: Path) -> list[dict]:
    """Extracts text from .docx Word documents using python-docx.

    Each paragraph group (~page equivalent) is returned as a numbered section.
    Since Word docs have no fixed pages, we split into chunks of 50 paragraphs.
    """

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Tables often contain the most useful structured content in Word docs.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    # Group into pseudo-pages of 50 paragraphs each
    page_size = 50
    pages = []
    for i in range(0, max(len(paragraphs), 1), page_size):
        text = _normalize_text("\n".join(paragraphs[i : i + page_size]))
        if text:
            pages.append(
                {
                    "page_number": (i // page_size) + 1,
                    "text": text,
                    "char_count": len(text),
                }
            )

    logger.info(
        f"python-docx extracted {len(pages)} sections from {file_path.name}"
    )
    return pages


def extract_text_pymupdf(pdf_path: Path) -> list[dict]:
    """Primary extractor using PyMuPDF.

    Returns list of {page_number, text, char_count}.
    PyMuPDF is ~10x faster than pypdf and handles most PDFs well.
    """

    pages = []
    doc = fitz.open(str(pdf_path))
    for page_num, page in enumerate(doc, start=1):
        text = _normalize_text(page.get_text("text"))
        pages.append(
            {
                "page_number": page_num,
                "text": text,
                "char_count": len(text),
            }
        )
    doc.close()
    logger.info(f"PyMuPDF extracted {len(pages)} pages from {pdf_path.name}")
    return pages


def extract_text_pdfplumber(pdf_path: Path) -> list[dict]:
    """Fallback extractor using pdfplumber.

    Better for PDFs with tables or complex column layouts.
    """

    pages = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = _normalize_text(page.extract_text() or "")
            pages.append(
                {
                    "page_number": page_num,
                    "text": text,
                    "char_count": len(text),
                }
            )
    logger.info(
        f"pdfplumber extracted {len(pages)} pages from {pdf_path.name}"
    )
    return pages


def extract_text_pypdf(pdf_path: Path) -> list[dict]:
    """Last-resort extractor using pypdf (pure Python, no C deps)."""

    pages = []
    reader = PdfReader(str(pdf_path))
    for page_num, page in enumerate(reader.pages, start=1):
        text = _normalize_text(page.extract_text() or "")
        pages.append(
            {
                "page_number": page_num,
                "text": text,
                "char_count": len(text),
            }
        )
    logger.info(f"pypdf extracted {len(pages)} pages from {pdf_path.name}")
    return pages


def extract_document_text(file_path: Path) -> tuple[list[dict], int]:
    """Main entry point for document text extraction.

    Routes to the correct extractor based on file extension.
      - .docx → python-docx
      - plain-text files (.txt/.md/.rst/.csv)
      - .pdf  → PyMuPDF → pdfplumber → pypdf (with fallback)

    Returns:
        pages: list of page dicts with text and metadata
        page_count: total number of pages/sections
    """

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()

    # ── Plain text files ─────────────────────────────────────────────────────
    if ext in TEXT_EXTENSIONS:
        text = _normalize_text(
            file_path.read_text(encoding="utf-8", errors="ignore")
        )
        if not text:
            return [], 0
        pages = [{"page_number": 1, "text": text, "char_count": len(text)}]
        logger.info(f"Text extractor loaded 1 section from {file_path.name}")
        return pages, 1

    # ── Word document ─────────────────────────────────────────────────────────
    if ext == ".docx":
        pages = extract_text_docx(file_path)
        pages = [p for p in pages if p["char_count"] > 0]
        return pages, len(pages)

    # ── PDF (with 3-library fallback) ─────────────────────────────────────────
    if ext != ".pdf":
        raise ValueError(f"Unsupported file type: {file_path.name}")

    pages: Optional[list[dict]] = None

    try:
        pages = extract_text_pymupdf(file_path)
    except Exception as e:
        logger.warning(f"PyMuPDF failed ({e}), trying pdfplumber...")

    if not pages or all(p["char_count"] == 0 for p in pages):
        try:
            pages = extract_text_pdfplumber(file_path)
        except Exception as e:
            logger.warning(f"pdfplumber failed ({e}), trying pypdf...")

    if not pages or all(p["char_count"] == 0 for p in pages):
        pages = extract_text_pypdf(file_path)

    pages = [p for p in pages if p["char_count"] > 0]
    logger.info(f"Final extraction: {len(pages)} non-empty pages from {file_path.name}")
    return pages, len(pages)


# Keep backward-compatible alias
extract_pdf_text = extract_document_text

